import os
import re
from docx import Document
from docx.shared import Inches
import fitz

class ConvertPdfToWordService:
    def __init__(self):
        pass

    def convert_pdf_to_word(self, pdf_path, word_path):
        # Open the PDF file
        pdf_file = fitz.open(pdf_path)

        # Create the output directory for images
        os.makedirs("images", exist_ok=True)

        # Create a new DOCX document
        doc = Document()

        for page_index in range(0, len(pdf_file)):
            # Load the page
            page = pdf_file.load_page(page_index)

            # Extract images from the page
            image_list = page.get_images(full=True)
            if image_list:
                for img_index, img in enumerate(image_list):
                    # Get the XREF of the image
                    xref = img[0]

                    # Extract the image bytes
                    base_image = pdf_file.extract_image(xref)
                    image_bytes = base_image["image"]

                    # Get the image extension
                    image_ext = base_image["ext"]

                    # Save the image
                    image_name = f"./images/image_{page_index + 1}_{img_index + 1}.{image_ext}"
                    with open(image_name, "wb") as image_file:
                        image_file.write(image_bytes)

                    # Add the image to the DOCX file
                    doc.add_picture(image_name, width=Inches(4))

            # Extract text from the page
            page_text = page.get_text("dict")
            if page_text and "blocks" in page_text:
                paragraph = ""
                for block in page_text["blocks"]:
                    if block.get("type") == 0:  # Text block
                        for line in block.get("lines", []):
                            line_text = "".join(span.get("text", "") for span in line.get("spans", []))
                            if line_text.strip():
                                # Accumulate lines into a paragraph until a sentence-ending punctuation is found
                                paragraph += (" " + line_text.strip()) if paragraph else line_text.strip()
                                
                    doc.add_paragraph(paragraph.strip())  # Add the paragraph text
                    paragraph = ""  # Reset for the next paragraph

                # Add any remaining text that didn't end with a punctuation
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

        # Save the DOCX file
        doc.save(word_path)
        pdf_file.close()
        print(f"Extraction completed. Text and images saved to {word_path}")