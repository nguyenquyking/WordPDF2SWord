from docx import Document

class DeleteFootnoteService:
    def __init__(self):
        pass
    
    def remove_footnotes_precisely(self, doc_path):
        """
        Remove footnotes and content after horizontal lines in a Word document,
        and save directly to the original document path.
        
        :param doc_path: Path to the Word document to be modified.
        """
        # Load the Word document
        doc = Document(doc_path)

        # Create a new document to save the cleaned content
        new_doc = Document()

        # Track whether we have reached the section containing footnotes
        in_footnotes = False

        for paragraph in doc.paragraphs:
            # If a paragraph contains a horizontal line, we enter the footnotes section
            if "__________" in paragraph.text:
                in_footnotes = True

            # If not in the footnotes section, copy the paragraph to the new document
            if not in_footnotes:
                new_doc.add_paragraph(paragraph.text, style=paragraph.style)

        # Overwrite the original document
        new_doc.save(doc_path)
        print(f"Footnotes and content after horizontal lines removed. Updated {doc_path}")